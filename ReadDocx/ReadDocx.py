from docx import Document
import re

from ReadDocx.ScheduleClass.ScheduleWeekGroup import ScheduleWeekGroup
from ReadDocx.Enum.TypeWeek import TypeWeek


class ReadDocx:
    schedule_list: list[ScheduleWeekGroup]

    def __init__(self, path_docx) -> None:
        self.path_docx = path_docx
        self.__document_to_list()

    def __document_to_list(self):
        document = Document(self.path_docx)
        table = document.tables

        schedule_group = []

        for i in range(len(table)):
            paragraph = document.paragraphs[i + 1].text

            number = re.findall(r'\bС\w*', paragraph)
            number = (number[0][1:])

            type_week = self.__get_type_week(paragraph)

            table_schedule = table[i]

            schedule_group.append(ScheduleWeekGroup(number, type_week, table_schedule))

        self.schedule_list = schedule_group

    def __get_type_week(self, paragraph_text: str):
        if TypeWeek.Denominator.value.name in paragraph_text:
            return TypeWeek.Denominator
        elif TypeWeek.Numerator.value.name in paragraph_text:
            return TypeWeek.Numerator
        else:
            raise Exception("Error -- get_type_week")
